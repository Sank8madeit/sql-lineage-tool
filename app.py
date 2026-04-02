import streamlit as st
import pandas as pd
from docx import Document
import sqlglot
from sqlglot import exp

# ---------------------------
# UI
# ---------------------------
st.set_page_config(page_title="SQL Lineage V4", layout="wide")
st.title("🚀 SQL Documentation Generator (FINAL V4)")

sql_input = st.text_area("Paste SQL Query", height=300)
generate = st.button("Generate Documentation")

# ---------------------------
# CORE ENGINE
# ---------------------------

class LineageEngine:

    def __init__(self, sql):
        self.sql = sql
        self.tree = None
        self.alias_map = {}
        self.cte_map = {}
        self.lineage = []

    # ---------------------------
    # Parse SQL safely
    # ---------------------------
    def parse_sql(self):
        try:
            self.tree = sqlglot.parse_one(self.sql, read="oracle")
        except Exception as e:
            raise Exception(f"SQL Parsing Failed: {e}")

    # ---------------------------
    # Extract alias → table
    # ---------------------------
    def build_alias_map(self):
        for table in self.tree.find_all(exp.Table):
            name = table.name
            alias = table.alias

            if alias:
                self.alias_map[alias] = name
            else:
                self.alias_map[name] = name

    # ---------------------------
    # Extract CTEs
    # ---------------------------
    def extract_ctes(self):
        for cte in self.tree.find_all(exp.CTE):
            self.cte_map[cte.alias] = cte.this

    # ---------------------------
    # Resolve table recursively
    # ---------------------------
    def resolve_table(self, table_name):

        # If direct table
        if table_name in self.alias_map:
            return self.alias_map[table_name]

        # If CTE → resolve inside
        if table_name in self.cte_map:
            cte_query = self.cte_map[table_name]

            for t in cte_query.find_all(exp.Table):
                return t.name  # return base table

        return table_name

    # ---------------------------
    # Process SELECT
    # ---------------------------
    def process_select(self, select):

        for proj in select.expressions:

            target = proj.alias_or_name
            expr = proj.this

            if not expr:
                continue

            cols = list(expr.find_all(exp.Column))

            if not cols:
                self.lineage.append({
                    "SOURCE_COLUMN": "N/A",
                    "SOURCE_TABLE": "N/A",
                    "TRANSFORMATION_LOGIC": expr.sql(),
                    "TARGET_COLUMN": target,
                    "TARGET_TABLE": "FINAL_OUTPUT"
                })

            else:
                for col in cols:
                    table_alias = col.table
                    real_table = self.resolve_table(table_alias)

                    self.lineage.append({
                        "SOURCE_COLUMN": col.name,
                        "SOURCE_TABLE": real_table,
                        "TRANSFORMATION_LOGIC": expr.sql(),
                        "TARGET_COLUMN": target,
                        "TARGET_TABLE": "FINAL_OUTPUT"
                    })

    # ---------------------------
    # Run Engine
    # ---------------------------
    def run(self):

        self.parse_sql()
        self.build_alias_map()
        self.extract_ctes()

        for node in self.tree.walk():
            if isinstance(node, exp.Select):
                self.process_select(node)

        return self.lineage


# ---------------------------
# DATA DICTIONARY
# ---------------------------

def generate_data_dict(df):
    df_unique = df.drop_duplicates(subset=["TARGET_COLUMN"])

    return pd.DataFrame({
        "Column Name": df_unique["TARGET_COLUMN"],
        "Description": df_unique["TRANSFORMATION_LOGIC"],
        "Type": "Derived"
    })


# ---------------------------
# DOC GENERATOR
# ---------------------------

def create_doc(df_sttm, df_dict):

    doc = Document()
    doc.add_heading("SQL Documentation V4", 0)

    doc.add_heading("STTM", 1)
    for _, row in df_sttm.iterrows():
        doc.add_paragraph(str(row.to_dict()))

    doc.add_heading("Data Dictionary", 1)
    for _, row in df_dict.iterrows():
        doc.add_paragraph(str(row.to_dict()))

    file = "SQL_V4_Documentation.docx"
    doc.save(file)
    return file


# ---------------------------
# MAIN
# ---------------------------

if generate:

    if not sql_input.strip():
        st.error("Please paste SQL")
    else:
        try:
            engine = LineageEngine(sql_input)
            lineage = engine.run()

            df = pd.DataFrame(lineage)

            df = df[[
                "SOURCE_COLUMN",
                "SOURCE_TABLE",
                "TRANSFORMATION_LOGIC",
                "TARGET_COLUMN",
                "TARGET_TABLE"
            ]]

            df_dict = generate_data_dict(df)

            # Excel
            excel_file = "SQL_V4_Output.xlsx"
            with pd.ExcelWriter(excel_file) as writer:
                df.to_excel(writer, sheet_name="STTM", index=False)
                df_dict.to_excel(writer, sheet_name="DataDictionary", index=False)

            # Doc
            doc_file = create_doc(df, df_dict)

            st.success("✅ FINAL V4 Generated Successfully")

            st.subheader("📊 STTM Preview")
            st.dataframe(df)

            st.download_button("Download Excel", open(excel_file, "rb"))
            st.download_button("Download DOCX", open(doc_file, "rb"))

        except Exception as e:
            st.error(str(e))